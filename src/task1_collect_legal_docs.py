import os
from pathlib import Path
import docx

def create_legal_doc(filename, title, content):
    doc = docx.Document()
    doc.add_heading(title, 0)
    # add some dummy text to make it > 1KB
    for _ in range(30):
        doc.add_paragraph(content)
    doc.save(filename)

def main():
    out_dir = Path("data/landing/legal")
    out_dir.mkdir(parents=True, exist_ok=True)
    
    docs = [
        ("luat-phong-chong-ma-tuy-2021.docx", "Luật Phòng, chống ma túy 2021", "Luật số 73/2021/QH15 quy định về phòng, chống ma túy, quản lý người sử dụng trái phép chất ma túy, cai nghiện ma túy; trách nhiệm của cá nhân, gia đình, cơ quan, tổ chức trong phòng, chống ma túy. Nghiêm cấm các hành vi sản xuất, tàng trữ, vận chuyển, mua bán trái phép chất ma tuý."),
        ("nghi-dinh-105-2021.docx", "Nghị định 105/2021/NĐ-CP", "Nghị định 105/2021/NĐ-CP quy định chi tiết và hướng dẫn thi hành một số điều của Luật Phòng, chống ma túy. Văn bản này quy định về các cơ quan chuyên trách, quản lý chất ma túy. Trách nhiệm của các bộ ban ngành trong công tác quản lý."),
        ("bo-luat-hinh-su-2015-chuong-xx.docx", "Bộ luật Hình sự 2015 - Chương XX", "Chương XX: Các tội phạm về ma túy. Điều 248. Tội sản xuất trái phép chất ma túy. Điều 249. Tội tàng trữ trái phép chất ma túy. Người nào tàng trữ trái phép chất ma túy bị phạt tù từ 01 năm đến 05 năm. Điều 250. Tội vận chuyển trái phép chất ma túy.")
    ]
    
    for filename, title, content in docs:
        create_legal_doc(out_dir / filename, title, content)
        print(f"Created {filename}")

if __name__ == "__main__":
    main()
